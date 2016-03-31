# coding=utf-8

from __future__ import unicode_literals

from compressor.utils.decorators import cached_property
from datetime import datetime, timedelta
from docx import Document
import StringIO
import re
import xlsxwriter

from django.conf import settings
from django.contrib import messages
from django.contrib.syndication.views import Feed
from django.contrib.contenttypes.models import ContentType
from django.core.urlresolvers import reverse, reverse_lazy
from django.db import transaction
from django.db.models import Prefetch
from django.http.response import HttpResponseRedirect, JsonResponse, Http404, HttpResponse
from django.shortcuts import get_object_or_404, redirect
from django.utils import timezone
from django.utils.encoding import force_text
from django.utils.formats import date_format
from django.utils.html import strip_tags
from django.utils.translation import ugettext_lazy as _, ugettext
from django.views.generic.base import View, TemplateView
from django.views.generic.detail import DetailView, SingleObjectMixin
from django.views.generic.edit import CreateView, UpdateView, FormMixin, FormView, \
    DeleteView
from django.views.generic.list import ListView

from libs.ajaxy.responses import AjaxyInlineRedirectResponse, AjaxyRedirectResponse, \
    AjaxyReloadResponse
from libs.attachtor import views as attachtor
from libs.attachtor.models.models import UploadGroup, Upload
from libs.attachtor.utils import get_upload_group_id
from wkhtmltopdf.views import PDFTemplateView

from account.models import User
from conversation.models import Conversation
from okmoderation.utils import get_moderated_form_class
from otakantaa.utils import send_email, send_email_to_multiple_receivers
from survey.conf import config as survey_config
from survey.forms import SurveyFormset
from survey.models import Survey, SurveyElement, SurveyQuestion
from survey.utils import survey_formset_initial, get_submitter

from . import perms
from .forms import AttachmentUploadForm, CreateParticipationForm, \
    CreateSchemeForm, SchemeSearchForm, ExportTextForm, UploadMultipleAttachmentsForm, \
    EditSchemeAdminOwnersForm, EditSchemeOwnersForm, AcceptOwnerInvitationForm, \
    CloseSchemeForm, ExportExcelForm
from .models import Scheme, ParticipationDetails, SchemeOwner
from .survey_perms import CanAnswerSurvey


class CreateSchemeView(CreateView):
    model = Scheme
    form_class = CreateSchemeForm
    template_name = 'scheme/create_scheme_form.html'

    def get_form_kwargs(self):
        kwargs = super(CreateSchemeView, self).get_form_kwargs()
        kwargs['user'] = self.request.user
        return kwargs

    def get_initial(self):
        initial = {'owners': [self.request.user, ]}
        return initial

    @transaction.atomic()
    def form_valid(self, form):
        obj = form.save(commit=False)
        obj.creator = self.request.user
        obj.save()
        form.save_m2m()
        form.save_owners()

        # action processing
        # action_performed.send(sender=form.instance, created=True)

        return HttpResponseRedirect(obj.get_absolute_url())

    def form_invalid(self, form):
        messages.error(self.request, _("Täytä kaikki pakolliset kentät."))
        return super(CreateSchemeView, self).form_invalid(form)


class PreFetchedObjectMixIn(object):
    obj_kwarg = 'obj'

    def get_object(self):
        return self.kwargs[self.obj_kwarg]


class SchemeDetailView(PreFetchedObjectMixIn, DetailView):
    model = Scheme
    template_name = 'scheme/scheme_detail.html'


class CreateParticipationBaseView(CreateView):
    model = ParticipationDetails
    form_class = CreateParticipationForm

    def get_initial(self, **kwargs):
        initials = super(CreateParticipationBaseView, self).get_initial()
        initials['scheme'] = get_object_or_404(Scheme, pk=self.kwargs['scheme_id'])
        return initials

    def form_invalid(self, form):
        messages.error(self.request, _("Täytä kaikki pakolliset kentät."))
        return super(CreateParticipationBaseView, self).form_invalid(form)


class CreateConversationView(CreateParticipationBaseView):
    template_name = 'participation/create_conversation_form.html'

    @transaction.atomic()
    def form_valid(self, form):
        conversation = Conversation.objects.create()
        obj = form.save(commit=False)
        obj.content_object = conversation
        obj.save()

        # action processing
        # action_performed.send(sender=form.instance, created=True)

        return HttpResponseRedirect(obj.get_absolute_url())


class CreateSurveyView(CreateParticipationBaseView):
    template_name = "participation/create_survey_form.html"

    def form_valid(self, form):
        survey = Survey.objects.create(show_results=survey_config.show_results_default)
        obj = form.save(commit=False)
        obj.content_object = survey
        obj.save()
        return HttpResponseRedirect(obj.get_absolute_url())


class ContentPartialEditView(PreFetchedObjectMixIn, UpdateView):
    def get_form_class(self):
        klass = self.kwargs['form_class']

        if not perms.OwnsScheme(
            request=self.request,
            obj=self.get_object()
        ).is_authorized():
            # we are moderating another user's content, mix in ModReasoningMixIn
            return get_moderated_form_class(klass, self.request.user)
        return klass

    def get_template_names(self):
        return [
            self.kwargs['template_name'],
            'content/edit_base_form.html'
        ]

    def form_valid(self, form):
        form.save()
        return AjaxyInlineRedirectResponse(
            reverse('content:scheme_detail_%s' % self.kwargs['fragment'],
                    kwargs={'scheme_id': self.kwargs['scheme_id']}))


class SchemePartialEditView(ContentPartialEditView):
    pass


class SchemeOwnersEditView(ContentPartialEditView):
    def get_form_class(self):
        scheme = self.get_object()
        if scheme.written_as_organization():
            klass = EditSchemeAdminOwnersForm
        else:
            klass = EditSchemeOwnersForm

        if not perms.OwnsScheme(
            request=self.request,
            obj=self.get_object()
        ).is_authorized():
            # we are moderating another user's content, mix in ModReasoningMixIn
            return get_moderated_form_class(klass, self.request.user)
        return klass

    def get_form(self, form_class=None):
        klass = self.get_form_class()
        kwargs = self.get_form_kwargs()
        return klass(self.request.user, **kwargs)

    def form_valid(self, form):
        if form.has_changed():
            """
            msg = render_to_string('okmessages/messages/owners_changed.txt',
                                   {'scheme': form.instance})
            """
            receivers = form.instance.owners.\
                exclude(user_id=self.request.user.pk).real().unique_users()

            """
            send_message(
                subject=_("Hankkeen omistajia on muutettu"),
                receivers=list(receivers),
                message=msg,
                sender=self.request.user
            )
            """

            send_email_to_multiple_receivers(
                _("Hankkeen omistajia on muutettu"),
                'scheme/email/owners_changed.txt',
                {'scheme': form.instance},
                receivers,
            )

        return super(SchemeOwnersEditView, self).form_valid(form)


class PreFetchedParticipationObjectMixIn(object):
    def get_object(self):
        return get_object_or_404(ParticipationDetails,
                                 pk=self.kwargs['participation_detail_id'])


class ParticipationPartialEditView(PreFetchedParticipationObjectMixIn,
                                   ContentPartialEditView):
    def get_form_class(self):
        klass = self.kwargs['form_class']

        if not perms.OwnsParticipation(
            request=self.request,
            obj=self.get_object()
        ).is_authorized():
            # we are moderating another user's content, mix in ModReasoningMixIn
            return get_moderated_form_class(klass, self.request.user)
        return klass

    def form_valid(self, form):
        form.save()
        return AjaxyInlineRedirectResponse(
            reverse('content:participation:participation_detail_%s' %
                    self.kwargs['fragment'],
                    kwargs={'scheme_id': self.kwargs['scheme_id'],
                            'participation_detail_id':
                                self.kwargs['participation_detail_id']}))


class ParticipationExpirationDateEditView(ParticipationPartialEditView):
    def form_valid(self, form):
        # do not reload page in vain
        if form.cleaned_data['expiration_date'] == self.get_object().expiration_date:
            return super(ParticipationExpirationDateEditView, self).form_valid(form)
        form.save()
        return AjaxyReloadResponse()


class ParticipationDetailView(PreFetchedParticipationObjectMixIn, DetailView):
    model = ParticipationDetails

    def get_context_data(self, **kwargs):
        context = super(ParticipationDetailView, self).get_context_data(**kwargs)
        context['scheme'] = get_object_or_404(Scheme, pk=self.kwargs['scheme_id'])
        return context


class ConversationDetailView(ParticipationDetailView):
    template_name = 'participation/conversation_detail.html'


class SurveyDetailView(ParticipationDetailView):
    template_name = "participation/survey_detail.html"
    edit_mode = False

    def get_context_data(self, **kwargs):
        context = super(SurveyDetailView, self).get_context_data(**kwargs)

        survey = self.get_object().content_object
        submitter = get_submitter(self.request)
        initial = survey_formset_initial(survey, submitter)
        can_answer = CanAnswerSurvey(request=self.request, obj=survey).is_authorized()
        context["formset"] = SurveyFormset(survey=survey, initial=initial,
                                           disabled=can_answer is False)

        if survey.elements.exists():
            context["edit_mode"] = self.edit_mode
        else:
            context["edit_mode"] = True

        context["survey"] = survey
        context["show_results_choices"] = survey_config.show_results_choices
        return context


class ContentPartialDetailView(SchemeDetailView):
    def get_template_names(self):
        return [self.kwargs['template_name'], ]


class ParticipationPartialDetailView(ParticipationDetailView):
    model = ParticipationDetails

    def get_template_names(self):
        return [self.kwargs['template_name'], ]


class UploadAttachmentView(attachtor.UploadAttachmentView):
    form_class = AttachmentUploadForm

    def get_form_kwargs(self):
        kwargs = super(UploadAttachmentView, self).get_form_kwargs()
        kwargs.update({
            'uploader': self.request.user,
            'upload_group': UploadGroup.objects.filter(pk=self.kwargs['upload_group_id'])
                                               .first()
        })
        return kwargs

    def form_invalid(self, form):
        if '__all__' in form.errors:
            error = form.errors['__all__'][0]
        elif 'file' in form.errors:
            error = form.errors['file'][0]
        else:
            error = ugettext("Tiedoston lähetys epäonnistui.")
        return JsonResponse({'error': error})


class UploadMultipleAttachmentsBaseView(PreFetchedObjectMixIn, FormView):
    template_name = 'content/attachments_form.html'
    form_class = UploadMultipleAttachmentsForm

    def get_form_kwargs(self):
        kwargs = super(UploadMultipleAttachmentsBaseView, self).get_form_kwargs()
        obj = self.get_object()
        lookup_kwargs = {'content_type': ContentType.objects.get_for_model(obj),
                         'object_id': obj.pk}
        try:
            group = UploadGroup.objects.get(**lookup_kwargs)
        except UploadGroup.DoesNotExist:
            group = UploadGroup(**lookup_kwargs)
            group.pk = get_upload_group_id()
            group.save()
        kwargs.update({'instance': group, 'user': self.request.user})
        return kwargs

    def get_context_data(self, **kwargs):
        context = super(UploadMultipleAttachmentsBaseView, self) \
            .get_context_data(**kwargs)
        context['object'] = self.get_object()
        return context

    def get_success_url(self):
        raise NotImplementedError()

    def form_valid(self, form):
        form.save(form)
        return self.get_success_url()

    def form_invalid(self, form):
        return self.render_to_response(self.get_context_data(form=form), status=400)


class UploadMultipleAttachmentsSchemeView(UploadMultipleAttachmentsBaseView):
    def get_success_url(self):
        return AjaxyInlineRedirectResponse(
            reverse('content:attachment_list', kwargs={
                'scheme_id': self.get_object().pk,
            }))


class UploadMultipleAttachmentsParticipationView(UploadMultipleAttachmentsBaseView):
    def get_success_url(self):
        obj = self.get_object()
        return AjaxyInlineRedirectResponse(
            reverse('content:participation:attachment_list', kwargs={
                'scheme_id': obj.scheme.pk,
                'participation_detail_id': obj.pk,
            }))


class AttachmentsDetailView(PreFetchedObjectMixIn, DetailView):
    template_name = 'content/attachments_list.html'


class DeleteAttachmentView(View):
    def delete(self, request, **kwargs):
        obj = get_object_or_404(Upload, pk=kwargs['upload_id'])
        obj.delete()
        return HttpResponse(content='')

    def post(self, request, **kwargs):
        return self.delete(request, **kwargs)


class DeleteSchemePictureView(PreFetchedObjectMixIn, View):
    def delete(self, request, **kwargs):
        obj = self.get_object()
        obj.picture.delete()
        obj.picture_alt_text = ''
        obj.save()
        return AjaxyInlineRedirectResponse(reverse('content:scheme_detail_picture',
                                                   kwargs={'scheme_id': obj.pk}))

    def post(self, request, **kwargs):
        return self.delete(request, **kwargs)


class SchemesView(FormMixin, ListView):
    http_method_names = ["get"]
    form_class = SchemeSearchForm
    object_list = None
    paginate_by = 30
    queryset = Scheme.objects \
        .filter(visibility=Scheme.VISIBILITY_PUBLIC) \
        .order_by("-published")
    template_name = "scheme/schemes.html"

    def get_context_data(self, **kwargs):
        context = super(SchemesView, self).get_context_data(**kwargs)

        page = int(self.request.GET.get("page", 1))
        paginator = context["paginator"]
        object_list = []
        for current_page in range(1, page + 1):
            object_list.extend(paginator.page(current_page).object_list)
        context["object_list"] = object_list

        context['rss_url'] = self.request.build_absolute_uri(reverse('content:rss'))
        return context

    def form_valid(self, form):
        schemes_display_template = None
        display_type = form.cleaned_data.get("display_type")
        if display_type:
            schemes_display_template = "scheme/scheme_{}.html".format(display_type)

        object_list = form.filtrate(self.get_queryset())

        context = self.get_context_data(
            form=form,
            schemes_display_template=schemes_display_template,
            object_list=object_list
        )
        return self.render_to_response(context)

    def get_form_kwargs(self):
        kwargs = super(SchemesView, self).get_form_kwargs()
        kwargs["data"] = self.request.GET if len(self.request.GET) else None
        kwargs["user"] = self.request.user
        return kwargs

    def get(self, request, *args, **kwargs):
        form_class = self.get_form_class()
        form = self.get_form(form_class)

        if len(request.GET):
            if form.is_valid():
                return self.form_valid(form)
            else:
                return self.form_invalid(form)

        object_list = self.get_queryset()
        context = self.get_context_data(form=form, object_list=object_list)
        return self.render_to_response(context)


class LoadSchemesView(SchemesView):
    template_name = None
    template_names = {
        SchemeSearchForm.DISPLAY_TYPE_BOXES: "scheme/scheme_boxes_results.html",
        SchemeSearchForm.DISPLAY_TYPE_LIST: "scheme/scheme_list.html"
    }

    def form_valid(self, form):
        display_type = form.cleaned_data["display_type"]
        if not display_type:
            display_type = form.DEFAULT_DISPLAY_TYPE
        self.template_name = self.template_names[display_type]
        object_list = form.filtrate(self.get_queryset())
        context = self.get_context_data(form=form, object_list=object_list)
        return self.render_to_response(context)

    def render_to_response(self, context, **response_kwargs):
        if not len(self.request.GET):
            self.template_name = \
                self.template_names[SchemeSearchForm.DEFAULT_DISPLAY_TYPE]
        return super(LoadSchemesView, self).render_to_response(context, **response_kwargs)


class SchemeStatusBaseView(View):
    def publish_scheme_and_participations(self, request, obj, set_message=True):
        obj.status = Scheme.STATUS_PUBLISHED
        obj.visibility = Scheme.VISIBILITY_PUBLIC
        obj.published = timezone.now()
        obj.save()
        for p in obj.participations.ready_to_publish():
            p.status = ParticipationDetails.STATUS_PUBLISHED
            p.save()

        if set_message:
            messages.success(request, ugettext("Hanke sekä siihen liittyvät "
                                               "kyselyt ja keskustelut on julkaistu!"))
        return obj

    def close_scheme_and_participations(self, request, obj, set_message=True):
        obj.status = Scheme.STATUS_CLOSED
        obj.closed = timezone.now()
        obj.save()
        for p in obj.participations.ready_to_close():
            p.expiration_date = datetime.today().date() - timedelta(1)
            p.save()

        if set_message:
            messages.success(request, ugettext("Hanke sekä siihen liittyvät "
                                               "kyselyt ja keskustelut on päätetty."))
        return obj


class PublishSchemeView(PreFetchedObjectMixIn, SchemeStatusBaseView):
    def post(self, request, **kwargs):
        scheme = self.get_object()
        self.publish_scheme_and_participations(request, scheme)
        return HttpResponseRedirect(reverse('content:scheme_detail',
                                            kwargs={'scheme_id': scheme.pk}))


class ReOpenSchemeView(View, PreFetchedObjectMixIn):
    def post(self, request, **kwargs):
        scheme = self.get_object()
        scheme.status = Scheme.STATUS_PUBLISHED
        scheme.closed = None
        scheme.save()
        messages.success(request, ugettext("Hanke on avattu uudelleen"))
        return AjaxyRedirectResponse(reverse('content:scheme_detail',
                                             kwargs={'scheme_id': scheme.pk}))


class RevertSchemeToDraftView(View, PreFetchedObjectMixIn):
    def post(self, request, **kwargs):
        scheme = self.get_object()
        scheme.status = Scheme.STATUS_DRAFT
        scheme.visibility = Scheme.VISIBILITY_DRAFT
        scheme.save()
        messages.success(request, ugettext("Hanke on palautettu luonnostilaan"))
        return AjaxyRedirectResponse(reverse('content:scheme_detail',
                                             kwargs={'scheme_id': scheme.pk}))


# todo: change to updateView when needed
class TryPublishSchemeView(PreFetchedObjectMixIn, DetailView, SchemeStatusBaseView):
    template_name = 'scheme/scheme_publish_preview.html'

    def get(self, request, *args, **kwargs):
        obj = self.get_object()
        if not obj.participations.ready_to_publish():
            self.publish_scheme_and_participations(request, obj)
            return redirect(reverse('content:scheme_detail',
                                    kwargs={'scheme_id': obj.pk}))
        return super(TryPublishSchemeView, self).get(request, *args, **kwargs)


class TryCloseSchemeView(PreFetchedObjectMixIn, UpdateView, SchemeStatusBaseView):
    template_name = 'scheme/scheme_close_preview.html'
    form_class = CloseSchemeForm

    def get(self, request, *args, **kwargs):
        # obj = self.get_object()
        # if not obj.participations.ready_to_close() and obj.closed:
        #    self.close_scheme_and_participations(request, obj)
        #    return redirect(self.get_success_url())
        return super(TryCloseSchemeView, self).get(request, *args, **kwargs)

    def form_valid(self, form):
        self.close_scheme_and_participations(self.request, self.get_object())
        return super(TryCloseSchemeView, self).form_valid(form)

    def get_success_url(self):
        return reverse('content:scheme_detail',
                       kwargs={'scheme_id': self.get_object().pk})


class UpdateSurveyShowResults(SingleObjectMixin, TemplateView):
    model = ParticipationDetails
    pk_url_kwarg = "participation_detail_id"
    template_name = "participation/survey_show_results_items.html"
    object = None

    def get_context_data(self, **kwargs):
        context = super(UpdateSurveyShowResults, self).get_context_data(**kwargs)
        context["show_results_choices"] = survey_config.show_results_choices
        return context

    def post(self, request, *args, **kwargs):
        self.object = self.get_object()
        survey = self.object.content_object
        value = int(self.kwargs["value"])
        assert value in dict(survey_config.show_results_choices).keys(), \
            "Invalid survey show_results value."
        survey.show_results = value
        survey.save()
        return self.render_to_response(self.get_context_data(survey=survey))


class ChangeParticipationStatusBaseView(SchemeStatusBaseView):
    STATUS_PUBLISHED = 'open'
    STATUS_DRAFT = 'draft'

    set_status = None

    def get_success_url(self, obj):
        url_part = 'content:participation:{}_detail'
        url_name = url_part.format('survey' if obj.is_survey() else 'conversation')
        return HttpResponseRedirect(
            reverse(url_name, kwargs={'scheme_id': obj.scheme.pk,
                                      'participation_detail_id': obj.pk}))

    def get_object(self):
        return get_object_or_404(ParticipationDetails,
                                 pk=self.kwargs['participation_detail_id'])

    def change_status(self, update_timestamp=True):
        pd_obj = self.get_object()
        pd_obj.status = self.set_status
        if update_timestamp:
            pd_obj.set_status_timestamp()
        pd_obj.save()
        return pd_obj

    def post(self, request, **kwargs):
        obj = self.change_status()
        messages.success(request, obj.label()[self.get_status_label()])
        return self.get_success_url(obj)

    def get_status_label(self):
        statuses = {
            ParticipationDetails.STATUS_PUBLISHED: self.STATUS_PUBLISHED,
            ParticipationDetails.STATUS_DRAFT: self.STATUS_DRAFT,
        }
        if self.set_status in statuses:
            return statuses[self.set_status]
        else:
            raise Http404("Wrong type for ParticipationDetail status")


class OpenParticipationView(ChangeParticipationStatusBaseView):
    set_status = ParticipationDetails.STATUS_PUBLISHED

    def post(self, request, **kwargs):
        obj = self.get_object()
        if obj.scheme.status == Scheme.STATUS_DRAFT:
            return HttpResponseRedirect(reverse('content:try_publish',
                                                kwargs={'scheme_id': obj.scheme.pk}))
        else:
            super(OpenParticipationView, self).post(request, **kwargs)

        return self.get_success_url(obj)


class EndParticipationView(ChangeParticipationStatusBaseView):
    def post(self, request, **kwargs):
        obj = self.get_object()
        obj.expiration_date = datetime.today().date() - timedelta(1)
        obj.save()
        return self.get_success_url(obj)


class SchemePdfExportView(PDFTemplateView):
    form_template_name = "scheme/scheme_export_form.html"
    template_name = "scheme/scheme_pdf.html"
    scheme = None

    def get_context_data(self, **kwargs):
        context = super(SchemePdfExportView, self).get_context_data(**kwargs)
        context["scheme"] = self.scheme
        context["format"] = "PDF"
        context["request"] = self.request
        return context

    def get_filename(self):
        return "otakantaa_{}.pdf".format(timezone.now().date())

    def get_form_kwargs(self):
        kwargs = {"scheme_id": self.kwargs["scheme_id"]}
        if len(self.request.GET):
            kwargs["data"] = self.request.GET
        return kwargs

    def get(self, request, *args, **kwargs):
        self.scheme = get_object_or_404(Scheme, pk=self.kwargs["scheme_id"])
        form = ExportTextForm(**self.get_form_kwargs())

        if len(request.GET) and form.is_valid():
            kwargs["scheme_information"] = form.cleaned_data["scheme_information"]
            if "conversations" in form.cleaned_data:
                kwargs["conversations"] = form.cleaned_data["conversations"] \
                    .prefetch_related("content_object__comments__votes",
                                      "content_object__comments__user",
                                      "content_object__comments__responses__votes",
                                      "content_object__comments__responses__user")
            if "surveys" in form.cleaned_data:
                kwargs["surveys"] = form.cleaned_data["surveys"].prefetch_related(
                    Prefetch("content_object__elements",
                             queryset=SurveyElement.objects.instance_of(SurveyQuestion),
                             to_attr="questions"),
                    "content_object__questions__answers",
                    "content_object__questions__options__answers",
                    "content_object__questions__survey__submissions"
                )
            emails = form.cleaned_data.get("emails")
            if emails:
                pdf_response = super(SchemePdfExportView, self) \
                    .get(request, *args, **kwargs)
                pdf_attachment = (self.get_filename(), pdf_response.render().content)
                for email in emails:
                    send_email(
                        ugettext("Otakantaa.fi hanke: %(title)s") %
                        {"title": force_text(self.scheme.title)},
                        "scheme/email/scheme_export_email.txt",
                        {"title": self.scheme.title, "format": "PDF"},
                        [email],
                        self.request.user.settings.language,
                        [pdf_attachment]
                    )
                messages.success(
                    request,
                    ugettext("%(format)s-tiedosto lähetetty liitetiedostona.") %
                    {"format": "PDF"}
                )
                return redirect("content:scheme_detail",
                                scheme_id=self.kwargs["scheme_id"])
            else:
                return super(SchemePdfExportView, self).get(request, *args, **kwargs)

        self.template_name = self.form_template_name
        self.response_class = self.html_response_class
        return TemplateView.render_to_response(
            self, context=self.get_context_data(form=form)
        )


class ExcelRenderer(object):
    columns = ()
    column_width = 10
    filename = "excel.xlsx"

    def get_filename(self):
        return self.filename

    def get_queryset(self):
        raise NotImplementedError()

    def get_columns(self):
        return self.columns

    @cached_property
    def format_bold(self):
        return self.workbook.add_format({'bold': 1})

    @cached_property
    def format_bolder(self):
        return self.workbook.add_format({'bold': 3})

    @cached_property
    def format_datetime(self):
        return self.workbook.add_format({'num_format': 'dd.mm.yyyy hh:mm:ss'})

    def write_footer_rows(self, worksheet, start_row):
        pass

    def render_response(self):
        buff = StringIO.StringIO()
        self.workbook = xlsxwriter.Workbook(buff)
        worksheet = self.workbook.add_worksheet()
        row, col = 0, 0
        columns = self.get_columns()

        for column in columns:
            worksheet.write(row, col, column['label'], self.get_column_format(row, col))
            worksheet.set_column(col, col, column.get('width', self.column_width))
            col += 1

        for obj in self.get_queryset():
            row += 1
            col = 0
            for column in self.get_columns():
                if column.get('formula', False):
                    method = worksheet.write_formula
                else:
                    method = worksheet.write
                method(row, col,
                       self.get_column_value(row, col, obj),
                       self.get_column_format(row, col, obj))
                col += 1
            self.write_footer_rows(worksheet, row + 1)

        self.workbook.close()
        buff.seek(0)
        resp = HttpResponse(buff.read(), content_type='application/octet-stream')
        resp["Content-Disposition"] = "attachment; filename={}".format(
            self.get_filename()
        )
        return resp

    def get_column_value(self, row, col, obj):
        column = self.get_columns()[col]
        if hasattr(self, 'get_%s' % column['attr']):
            return getattr(self, 'get_%s' % column['attr'])(obj, row, col)
        v = getattr(obj, column['attr'], None)
        if callable(v):
            v = v()
        return v

    def get_column_format(self, row, col, obj=None):
        column = self.get_columns()[col]
        if row == 0:
            if 'header_format' in column:
                return column['header_format']
            else:
                return self.format_bold
        else:
            if 'body_format' in column:
                return column['body_format']
            else:
                return None


class ParticipationExcelRenderer(ExcelRenderer):
    def get_filename(self):
        return "otakantaa_{}.xlsx".format(timezone.now().date())

    def __init__(self, form=None, *args, **kwargs):
        self.form = form
        super(ParticipationExcelRenderer, self).__init__(*args, **kwargs)


class ConversationExcelRenderer(ParticipationExcelRenderer):
    columns = (
        {"attr": "id", "label": "ID"},
        {"attr": "target_comment_id", "label": "Parent ID"},
        {"attr": "title", "label": ugettext("Otsikko"), "width": 20},
        {"attr": "comment", "label": ugettext("Teksti"), "width": 20},
        {"attr": "votes_up", "label": ugettext("Peukku ylös")},
        {"attr": "votes_down", "label": ugettext("Peukku alas")},
        {"attr": "votes_sum", "label": ugettext("Peukkujen summa")},
        {"attr": "user", "label": ugettext("Tekijä")},
        {"attr": "language", "label": ugettext("Kieli")},
        {"attr": "municipality", "label": ugettext("Kotikunta")},
        {"attr": "organization", "label": ugettext("Organisaatio")},
    )

    def get_queryset(self):
        conversation = self.form.cleaned_data["conversation"]
        queryset = conversation.content_object.comments \
            .visible() \
            .select_related("user__settings__municipality", "target_comment") \
            .prefetch_related("user__organizations", "votes")
        return queryset

    def get_columns(self):
        columns = list(self.columns)
        columns.append({
            "attr": "created",
            "label": ugettext("Aika"),
            "body_format": self.format_datetime
        })
        return columns

    def get_column_value(self, row, col, obj):
        column = self.get_columns()[col]
        if column["attr"] == "user":
            return force_text(obj.user) if obj.user is not None else obj.user_name
        elif column["attr"] == "language":
            if obj.user:
                return obj.user.settings.language
            return ""
        elif column["attr"] == "votes_up":
            return sum([vote.choice for vote in obj.votes.all()
                        if vote.choice == vote.VOTE_UP])
        elif column["attr"] == "votes_down":
            return sum([vote.choice for vote in obj.votes.all()
                        if vote.choice == vote.VOTE_DOWN])
        elif column["attr"] == "votes_sum":
            return sum([vote.choice for vote in obj.votes.all()])
        elif column["attr"] == "municipality":
            if obj.user:
                return obj.user.settings.municipality.name
            return ""
        elif column["attr"] == "organization":
            if obj.user:
                return ", ".join(
                    ["{}".format(org.name) for org in obj.user.organizations.all()]
                )
            return ""
        elif column["attr"] == "created":
            return timezone.make_naive(obj.created)
        return super(ConversationExcelRenderer, self).get_column_value(row, col, obj)


class SurveyExcelRenderer(ParticipationExcelRenderer):
    def __init__(self, *args, **kwargs):
        super(SurveyExcelRenderer, self).__init__(*args, **kwargs)
        self.participation_details = self.form.cleaned_data["survey"]
        self.survey = self.participation_details.content_object

    def get_queryset(self):
        return self.survey.submissions \
            .select_related("survey", "submitter__user__settings__municipality") \
            .prefetch_related("submitter__user__organizations", "answers__option")

    def get_columns(self):
        columns = [
            {"attr": "created", "label": ugettext("Vastausaikaika"),
             "body_format": self.format_datetime},
            {"attr": "user", "label": ugettext("Vastaaja (Nimimerkki)")},
            {"attr": "municipality", "label": ugettext("Kotikunta")},
            {"attr": "language", "label": ugettext("Kieli")},
            {"attr": "organization", "label": ugettext("Organisaatio")},
        ]

        for question in self.survey.questions:
            columns.append({
                "attr": "question_{}".format(question.pk),
                "label": force_text(question.text)
            })

        return columns

    def get_column_value(self, row, col, obj):
        column = self.get_columns()[col]
        if column["attr"] == "user":
            if obj.submitter.user:
                return obj.submitter.user.get_full_name()
            return ""
        elif column["attr"] == "language":
            if obj.submitter.user:
                return obj.submitter.user.settings.language
            return ""
        elif column["attr"] == "municipality":
            if obj.submitter.user:
                return obj.submitter.user.settings.municipality.name
            return ""
        elif column["attr"] == "organization":
            if obj.submitter.user:
                return ", ".join(
                    ["{}".format(org.name)
                     for org in obj.submitter.user.organizations.all()]
                )
            return ""
        elif column["attr"] == "created":
            return timezone.make_naive(obj.created)
        elif column["attr"].startswith("question_"):
            return self.get_question_column_value(column, obj)
        return super(SurveyExcelRenderer, self).get_column_value(row, col, obj)

    def get_question(self, question_id):
        for question in self.survey.questions:
            if question.pk == question_id:
                return question

    def get_question_column_value(self, column, obj):
        question_id = int(re.search(r"question_(?P<pk>\d+)$", column["attr"]).group("pk"))
        question = self.get_question(question_id)
        answers = [answer for answer in obj.answers.all()
                   if answer.question_id == question_id]
        if question.type == question.TYPE_CHECKBOX:
            if answers:
                return ", ".join([force_text(answer.option.text) for answer in answers
                                  if answer.option])
            else:
                return ""
        elif question.type == question.TYPE_RADIO:
            answer = answers[0]
            if answer.option:
                return force_text(answer.option.text)
            else:
                return ""
        else:
            answer = answers[0]
            return answer.text


class ParticipationExcelExportView(TemplateView):
    template_name = "scheme/scheme_export_excel_form.html"
    renderer_map = {
        Conversation: ConversationExcelRenderer,
        Survey: SurveyExcelRenderer,
    }
    scheme = None

    def get_context_data(self, **kwargs):
        context = super(ParticipationExcelExportView, self).get_context_data(**kwargs)
        context["scheme"] = self.scheme
        context["format"] = "Excel"
        return context

    def get_form_kwargs(self):
        kwargs = {"scheme_id": self.kwargs["scheme_id"]}
        if len(self.request.GET):
            kwargs["data"] = self.request.GET
        return kwargs

    def get(self, request, *args, **kwargs):
        self.scheme = get_object_or_404(Scheme, pk=self.kwargs["scheme_id"])
        form = ExportExcelForm(**self.get_form_kwargs())

        if len(request.GET) and form.is_valid():
            renderer_class = self.renderer_map[type(form.get_participation())]
            renderer = renderer_class(form=form)

            emails = form.cleaned_data.get("emails")
            if emails:
                excel_attachment = (renderer.get_filename(),
                                    renderer.render_response().content)
                for email in emails:
                    send_email(
                        ugettext("Otakantaa.fi hanke: %(title)s") %
                        {"title": force_text(self.scheme.title)},
                        "scheme/email/scheme_export_email.txt",
                        {"title": self.scheme.title, "format": "Excel"},
                        [email],
                        self.request.user.settings.language,
                        [excel_attachment]
                    )
                messages.success(
                    request,
                    ugettext("%(format)s-tiedosto lähetetty liitetiedostona.") %
                    {"format": "Excel"}
                )
                return redirect("content:scheme_detail",
                                scheme_id=self.kwargs["scheme_id"])
            else:
                return renderer.render_response()

        return self.render_to_response(self.get_context_data(form=form))


class SchemeWordExportView(TemplateView):
    template_name = "scheme/scheme_export_form.html"
    scheme = None

    def get_context_data(self, **kwargs):
        context = super(SchemeWordExportView, self).get_context_data(**kwargs)
        context["scheme"] = self.scheme
        context["format"] = "Word"
        return context

    def get_filename(self):
        return "otakantaa_{}.docx".format(timezone.now().date())

    def get_form_kwargs(self):
        kwargs = {"scheme_id": self.kwargs["scheme_id"]}
        if len(self.request.GET):
            kwargs["data"] = self.request.GET
        return kwargs

    def set_document_comment(self, document, comment):
        document.add_paragraph(
            "@{} +{} -{} ={} ({}) ".format(
                comment.user_name,
                comment.votes.up().count(),
                comment.votes.down().count(),
                comment.votes.balance(),
                date_format(comment.created, "SHORT_DATETIME_FORMAT")
            )
        )
        document.add_paragraph(comment.comment)

    def set_document_conversations(self, document, conversations_details):
        document.add_heading(ugettext("Keskustelut"), 1)

        for conversation_details in conversations_details:
            conversation = conversation_details.content_object
            document.add_heading(
                strip_tags(force_text(conversation_details.title)), 2
            )
            document.add_paragraph(
                strip_tags(force_text(conversation_details.description))
            )

            for comment in conversation.comments.comments().visible():
                document.add_heading(ugettext("Kommentti"), 3)
                self.set_document_comment(document, comment)

                responses = comment.responses.all()
                if responses:
                    document.add_heading(ugettext("Vastaukset"), 4)

                    for response in responses:
                        self.set_document_comment(document, response)

    def set_document_survey_answers(self, document, question):
        document.add_heading(ugettext("Vastaukset"), 3)

        if question.type == question.TYPE_TEXT:
            answers = question.answers.all()
            if answers:
                for answer in answers:
                    document.add_paragraph(answer.text)
            else:
                document.add_paragraph().add_run(ugettext("Ei vastauksia")) \
                    .font.italic = True

        elif question.type == question.TYPE_NUMBER:
            number_answers = [answer.text for answer in question.answers.all()]
            if number_answers:
                document.add_paragraph(", ".join(number_answers))
            else:
                document.add_paragraph().add_run(ugettext("Ei vastauksia")) \
                    .font.italic = True

        elif question.type in (question.TYPE_RADIO, question.TYPE_CHECKBOX):
            submission_count = question.survey.submissions.count()
            for option in question.options.all():
                document.add_paragraph(
                    "{} {}/{}".format(
                        option.text,
                        option.answers.count(),
                        submission_count
                    )
                )

    def set_document_surveys(self, document, surveys_details):
        document.add_heading(ugettext("Kyselyt"), 1)

        for survey_details in surveys_details:
            survey = survey_details.content_object
            document.add_heading(strip_tags(force_text(survey_details.title)), 2)
            document.add_paragraph(strip_tags(force_text(survey_details.description)))

            for question in survey.elements.questions():
                document.add_heading(ugettext("Kysymys"), 3)
                if question.required:
                    document.add_paragraph("{} *".format(force_text(question.text)))
                else:
                    document.add_paragraph(force_text(question.text))
                if len(question.instruction_text) > 1:
                    instruction = document \
                        .add_paragraph() \
                        .add_run(force_text(question.instruction_text))
                    instruction.font.italic = True

                self.set_document_survey_answers(document, question)

    def get_document(self, form):
        document = Document()
        document.add_heading(force_text(self.scheme.title), 0)

        if form.cleaned_data["scheme_information"]:
            document.add_heading(ugettext("Tiedot"), 1)
            document.add_heading(ugettext("Yhteenveto"), 2)
            document.add_paragraph(strip_tags(force_text(self.scheme.summary)))
            document.add_heading(ugettext("Lisätiedot"), 2)
            document.add_paragraph(strip_tags(force_text(self.scheme.description)))

        if 'conversations' in form.cleaned_data:
            self.set_document_conversations(document, form.cleaned_data["conversations"])

        if 'surveys' in form.cleaned_data:
            self.set_document_surveys(document, form.cleaned_data["surveys"])

        return document

    def render_to_word_response(self, form):
        response = HttpResponse(content_type="text/docx")
        response["Content-Disposition"] = "attachment; filename={}" \
            .format(self.get_filename())
        document = self.get_document(form)
        document.save(response)
        return response

    def get(self, request, *args, **kwargs):
        self.scheme = get_object_or_404(Scheme, pk=self.kwargs["scheme_id"])
        form = ExportTextForm(**self.get_form_kwargs())

        if len(request.GET) and form.is_valid():
            emails = form.cleaned_data.get("emails")
            if emails:
                word_response = self.render_to_word_response(form)
                word_attachment = (self.get_filename(), word_response.content)
                for email in emails:
                    send_email(
                        ugettext("Otakantaa.fi hanke: %(title)s") %
                        {"title": force_text(self.scheme.title)},
                        "scheme/email/scheme_export_email.txt",
                        {"title": self.scheme.title, "format": "Word"},
                        [email],
                        self.request.user.settings.language,
                        [word_attachment]
                    )
                messages.success(
                    request,
                    ugettext("%(format)s-tiedosto lähetetty liitetiedostona.") %
                    {"format": "Word"}
                )
                return redirect("content:scheme_detail",
                                scheme_id=self.kwargs["scheme_id"])
            else:
                return self.render_to_word_response(form)

        return self.render_to_response(context=self.get_context_data(form=form))


class SchemeOwnerInvitationView(UpdateView):
    template_name = 'scheme/scheme_owner_invitation.html'
    form_class = AcceptOwnerInvitationForm
    model = SchemeOwner

    def get_object(self, queryset=None):
        return get_object_or_404(SchemeOwner, pk=self.kwargs.get('scheme_owner_id', None))

    def form_valid(self, form):
        obj = form.save(commit=False)
        if obj.approved:
            obj.save()
            messages.success(self.request, _("Sinut on lisätty hankkeen omistajaksi"))
        else:
            obj.delete()
            sender = obj.organization if obj.organization else obj.user
            options = {'obj': sender, 'scheme': obj.scheme}
            """
            msg = render_to_string('okmessages/messages/owner_invitation_denied.txt',
                                   options)
            """
            pk_list = obj.scheme.owners.real().values_list('user_id', flat=True)

            """
            send_message(
                subject=_("Kieltäytyminen omistajakutsusta"),
                receivers=list(User.objects.filter(pk__in=pk_list)),
                message=msg,
                sender=self.request.user
            )
            """

            send_email_to_multiple_receivers(
                _("Kieltäytyminen omistajakutsusta"),
                'scheme/email/owner_invitation_denied.txt',
                options,
                list(User.objects.filter(pk__in=pk_list)),
            )

            messages.success(self.request, _("Vastauksesi on käsitelty"))
        return HttpResponseRedirect(reverse('account:messages', kwargs={
            'user_id': obj.user.pk}))


class DeleteSchemeView(DeleteView):
    model = Scheme
    pk_url_kwarg = 'scheme_id'
    template_name = 'scheme/scheme_confirm_delete.html'
    success_url = reverse_lazy('schemes')

    def get_success_url(self):
        messages.success(
            self.request,
            ugettext("Hanke '{0}' on poistettu.").format(self.get_object().title)
        )
        return super(DeleteSchemeView, self).get_success_url()


class DeleteParticipationDetailView(DeleteView):
    model = ParticipationDetails
    pk_url_kwarg = 'participation_detail_id'
    template_name = 'participation/participation_confirm_delete.html'

    def get_success_url(self):
        messages.success(self.request, ugettext("Poistaminen onnistui."))
        return reverse('content:scheme_detail', kwargs={
            'scheme_id': self.get_object().scheme.pk})


class SchemeFeed(Feed):
    title = _("otakantaa.fi")
    description = _("Seuraa hankkeita")
    link = settings.BASE_URL
    form_class = SchemeSearchForm
    queryset = Scheme.objects.get_queryset()

    def get_form_kwargs(self):
        kwargs = super(SchemeFeed, self).get_form_kwargs()
        kwargs["data"] = self.request.GET if len(self.request.GET) else None
        kwargs["user"] = self.request.user
        return kwargs

    def get_object(self, request, *args, **kwargs):
        form = self.form_class(request.GET)
        if not form.is_valid():
            raise Http404()
        return form

    def items(self, form):
        qs = self.queryset.filter(visibility=Scheme.VISIBILITY_PUBLIC).\
            order_by('-published')
        return form.filtrate(qs)

    def item_title(self, item):
        return item

    def item_description(self, item):
        return strip_tags('%s' % item.lead_text)

    def item_link(self, item):
        return self.link.rstrip('/') + item.get_absolute_url()


class PremoderationToggleView(View):
    @transaction.atomic()
    def post(self, request, *args, **kwargs):
        obj = self.kwargs['obj']
        obj.premoderation = bool(int(kwargs['premoderation_state']))
        obj.save()
        if obj.premoderation:
            messages.success(
                self.request,
                ' '.join([
                    ugettext("Kommenttien esimoderointi on otettu käyttöön."),
                    ugettext("Hankkeeseen lisättävät kommentit menevät palvelun "
                             "moderaattorien hyväksyttäväksi ennen julkaisua.")
                ])
            )
        else:
            messages.success(
                self.request,
                ' '.join([
                    ugettext("Kommenttien esimoderointi on poistettu käytöstä."),
                    ugettext("Hankkeeseen lisättävät kommentit julkaistaan välittömästi.")
                ])
            )
        return AjaxyReloadResponse()


class SetInteractionView(SingleObjectMixin, TemplateView):
    model = Scheme
    pk_url_kwarg = 'scheme_id'
    template_name = 'scheme/scheme_interaction_items.html'
    object = None

    def post(self, request, *args, **kwargs):
        obj = self.kwargs['obj']
        obj.interaction = int(kwargs['interaction'])
        obj.save()
        return self.render_to_response(self.get_context_data(object=obj))


class SchemeTwitterSearchFormView(UpdateView):
    model = Scheme
    fields = ["twitter_search"]
    pk_url_kwarg = "scheme_id"
    template_name = "scheme/scheme_twitter_search_form.html"

    def form_valid(self, form):
        success_msg = "{} {}".format(
            ugettext("Twitter-syötteen hakusana vaihdettu."),
            ugettext("Päivittymisessä voi kestää jonkin aikaa.")
        )
        messages.success(self.request, success_msg)
        return super(SchemeTwitterSearchFormView, self).form_valid(form)
